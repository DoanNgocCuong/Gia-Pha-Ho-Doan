"""
========================================================================
 Công cụ 2: TRÍCH XUẤT VĂN BẢN + THÔNG TIN THỤT LỀ (CHI TIẾT)
========================================================================

Mục đích:
---------
- Đọc file Word .docx và xuất ra text kèm theo thông tin định dạng:
    + left (twips)  : độ thụt lề trái của paragraph
    + ilvl          : cấp (level) nếu paragraph thuộc danh sách numbered list
    + tabs          : số tab ký tự ở đầu dòng
- Đây là công cụ KHẢO SÁT chính - output của nó chính là "bản đồ thế hệ"
  mà ta dùng để xây cây gia phả ở bước build_tree.py.

Khi nào dùng:
-------------
- Khi bạn sửa/thêm nội dung trong file Word gốc và muốn biết Word đã
  ghi nhận thế hệ (indent) như thế nào.
- Khi cây gia phả hiện ra sai thế hệ - dùng file này để dò lỗi dòng nào
  bị lệch lề trong Word.

Đơn vị "twips" là gì?
---------------------
- Twip = 1/20 point = 1/1440 inch.
- Word lưu độ thụt lề ở đơn vị twip. Quy đổi dễ nhớ:
    720 twips  = 0.5 inch = 1 bước thụt lề tiêu chuẩn  -> tương ứng 1 thế hệ
    1440 twips = 1.0 inch = 2 bước (thế hệ)
    2160 twips = 1.5 inch = 3 bước (thế hệ)

Yêu cầu:
--------
- Python 3.8+
- Thư viện python-docx:  pip install python-docx

Cách chạy (PowerShell):
-----------------------
    python utils\extract_with_indent.py

Sẽ tạo ra file:
    d:\GIT\Gia-Pha-Ho-Doan\utils\_out_indent.txt

Định dạng file đầu ra (mỗi dòng):
    [<idx>] left=<twips|None> ilvl=<số|None> tabs=<n> | <nội dung>

Ví dụ:
    [004] left=2160 ilvl=None tabs=0 | ô Vĩnh

Đọc nhanh: dòng "ô Vĩnh" có thụt lề 2160 twips → thế hệ thứ 4 (theo
công thức ở build_tree.py: depth = left/720 + 1 = 4).

Nguyên tắc thiết kế (SOLID):
---------------------------
- Single Responsibility: file chỉ chịu trách nhiệm DÒ thông tin thụt lề.
- Open/Closed: hàm get_indent_info có thể mở rộng để đọc thêm thuộc tính
  khác (hanging, firstLine...) mà không cần sửa phần còn lại.
========================================================================
"""

from docx import Document
from docx.oxml.ns import qn  # qn = qualified name - dùng để truy cập thẻ XML có namespace

# ──────────────────────────────────────────────────────────────────────
# CẤU HÌNH ĐƯỜNG DẪN
# ──────────────────────────────────────────────────────────────────────
DOCX_PATH = r'd:\GIT\Gia-Pha-Ho-Doan\Cụ Liễu.docx'
OUT_PATH = r'd:\GIT\Gia-Pha-Ho-Doan\utils\_out_indent.txt'


def get_indent_info(paragraph) -> tuple:
    """
    Đọc thông tin thụt lề từ XML gốc của paragraph.

    Word lưu định dạng ở dạng XML. Đường dẫn ta quan tâm:
        w:p -> w:pPr -> w:ind (thuộc tính w:left / w:start)
        w:p -> w:pPr -> w:numPr -> w:ilvl

    Tham số:
        paragraph: đối tượng Paragraph của python-docx

    Trả về:
        (left_twips, ilvl)
        - left_twips : int hoặc None  (độ thụt lề tính bằng twips)
        - ilvl       : int hoặc None  (cấp trong numbered list, nếu có)
    """
    # Truy cập phần tử <w:pPr> - paragraph properties
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is None:
        # Paragraph không có khối properties -> không có thông tin thụt lề
        return None, None

    # ── 1. Đọc độ thụt lề trái (left) ──────────────────────────────
    left = None
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        # Ưu tiên w:left, fallback w:start (Word đời mới dùng w:start)
        for attr in ('w:left', 'w:start'):
            value = ind.get(qn(attr))
            if value is not None:
                try:
                    left = int(value)
                    break  # Đọc được một trong hai là đủ
                except ValueError:
                    # Giá trị không phải số nguyên -> bỏ qua
                    pass

    # ── 2. Đọc cấp (level) trong numbered list ─────────────────────
    ilvl = None
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        ilvl_elem = numPr.find(qn('w:ilvl'))
        if ilvl_elem is not None:
            value = ilvl_elem.get(qn('w:val'))
            if value is not None:
                try:
                    ilvl = int(value)
                except ValueError:
                    pass

    return left, ilvl


def count_leading_tabs(text: str) -> int:
    """Đếm số tab ở đầu chuỗi (giống file extract_text.py)."""
    tabs = 0
    for ch in text:
        if ch == '\t':
            tabs += 1
        else:
            break
    return tabs


def dump_paragraphs(docx_path: str, out_path: str) -> int:
    """
    Đọc file .docx và ghi ra file text kèm đầy đủ metadata thụt lề.

    Tham số:
        docx_path : đường dẫn file Word
        out_path  : đường dẫn file text đầu ra

    Trả về:
        int - tổng số paragraph đã xuất
    """
    doc = Document(docx_path)

    with open(out_path, 'w', encoding='utf-8') as f:
        for index, paragraph in enumerate(doc.paragraphs):
            raw_text = paragraph.text
            leading_tabs = count_leading_tabs(raw_text)
            left, ilvl = get_indent_info(paragraph)

            # Format mỗi dòng với index 3 chữ số (000, 001, ...) để sort dễ
            f.write(
                f'[{index:03d}] '
                f'left={left} '
                f'ilvl={ilvl} '
                f'tabs={leading_tabs} | '
                f'{raw_text}\n'
            )

    return len(doc.paragraphs)


def main() -> None:
    """Hàm điều phối: đọc docx -> xuất file metadata thụt lề."""
    print(f'[1/2] Đang đọc & phân tích: {DOCX_PATH}')
    total = dump_paragraphs(DOCX_PATH, OUT_PATH)

    print(f'[2/2] Đã ghi vào: {OUT_PATH}')
    print(f'Hoàn tất! Tổng {total} paragraph với metadata thụt lề.')
    print()
    print('Gợi ý đọc nhanh:')
    print('  - left=720   → thế hệ 2 (I. Cụ Hán / II. Cụ Quyết / III. Cụ Huấn)')
    print('  - left=1440  → thế hệ 3 (con/vợ của 3 cụ trên)')
    print('  - left=2160  → thế hệ 4 ...')
    print('  - left=N, tabs=k → thế hệ = (N/720 + 1) + k')


if __name__ == '__main__':
    main()
